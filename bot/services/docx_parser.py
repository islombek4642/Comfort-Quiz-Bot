"""
DOCX Parser Service
Word fayldan savollarni o'qib olish
"""
import re
import uuid
from dataclasses import dataclass
from typing import Optional
from docx import Document
from bot.models import Question


@dataclass
class ParseResult:
    """Parse natijasi"""
    success: bool
    questions: list[Question]
    error_message: str = ""
    warnings: list[str] = None
    
    def __post_init__(self):
        if self.warnings is None:
            self.warnings = []


class DocxParser:
    """
    DOCX fayllarni parse qilish.
    
    Qo'llab-quvvatlanadigan formatlar:
    
    Format 1 (Raqamli):
    1. Savol matni?
    A) Variant 1
    B) Variant 2
    *C) To'g'ri javob
    D) Variant 4
    
    Format 2 (So'roq belgisi bilan - YANGI):
    ?Savol matni
    +To'g'ri javob
    =Noto'g'ri variant
    =Noto'g'ri variant
    =Noto'g'ri variant
    
    Format 3 (Yulduzcha bilan):
    1. Savol matni?
    A) Variant 1
    *B) Variant 2
    C) Variant 3
    """
    
    # Savol boshlanish pattern'lari
    QUESTION_PATTERNS = [
        r'^\?+\s*(.+)$',  # ?Savol matni (yangi format)
        r'^(\d+)\s*[\.\)\-]\s*(.+)$',  # 1. Savol yoki 1) Savol
        r'^Savol\s*(\d+)\s*[\.\:\)]\s*(.+)$',  # Savol 1: Savol matni
        r'^S\s*(\d+)\s*[\.\:\)]\s*(.+)$',  # S1. Savol matni
    ]
    
    # Variant pattern'lari (klassik A, B, C, D)
    OPTION_PATTERNS = [
        r'^([A-Za-z])\s*[\)\.\-]\s*(.+)$',  # A) variant yoki A. variant
        r'^(\d+)\s*[\)\.\-]\s*(.+)$',  # 1) variant
    ]
    
    # To'g'ri javob belgilari
    CORRECT_MARKERS = ['*', '+', '✓', '✔', '√']
    
    def __init__(self):
        self.questions: list[Question] = []
        self.warnings: list[str] = []
    
    async def parse_file(self, file_path: str) -> ParseResult:
        """DOCX faylni parse qilish"""
        try:
            doc = Document(file_path)
            return await self._parse_document(doc)
        except Exception as e:
            return ParseResult(
                success=False,
                questions=[],
                error_message=f"Faylni o'qishda xato: {str(e)}"
            )
    
    async def parse_bytes(self, file_bytes) -> ParseResult:
        """Bytes dan parse qilish"""
        import io
        try:
            doc = Document(io.BytesIO(file_bytes))
            return await self._parse_document(doc)
        except Exception as e:
            return ParseResult(
                success=False,
                questions=[],
                error_message=f"Faylni o'qishda xato: {str(e)}"
            )
    
    async def _parse_document(self, doc: Document) -> ParseResult:
        """Document obyektini parse qilish"""
        self.questions = []
        self.warnings = []
        
        # Barcha paragraflarni olish
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        if not paragraphs:
            return ParseResult(
                success=False,
                questions=[],
                error_message="❌ Fayl bo'sh yoki matn topilmadi."
            )
        
        # Formatni aniqlash
        has_question_mark_format = any(p.startswith('?') for p in paragraphs)
        
        if has_question_mark_format:
            # Yangi format: ?savol, +to'g'ri, =noto'g'ri
            return await self._parse_question_mark_format(paragraphs)
        else:
            # Klassik format: 1. Savol, A) variant
            return await self._parse_classic_format(paragraphs)
    
    async def _parse_question_mark_format(self, paragraphs: list[str]) -> ParseResult:
        """
        Yangi format parse qilish:
        ?Savol matni
        +To'g'ri javob
        =Noto'g'ri variant
        """
        current_question = None
        current_options = []
        correct_index = -1
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Savol (? bilan boshlanadi)
            if para.startswith('?'):
                # Oldingi savolni saqlash
                if current_question and current_options:
                    self._save_question(current_question, current_options, correct_index)
                
                # Yangi savol - barcha ? belgilarini olib tashlash
                current_question = para.lstrip('?').strip()
                current_options = []
                correct_index = -1
                continue
            
            # To'g'ri javob (+ bilan boshlanadi)
            if para.startswith('+') and current_question:
                option_text = para[1:].strip()
                current_options.append(option_text)
                correct_index = len(current_options) - 1
                continue
            
            # Noto'g'ri variant (= bilan boshlanadi)
            if para.startswith('=') and current_question:
                option_text = para[1:].strip()
                current_options.append(option_text)
                continue
        
        # Oxirgi savolni saqlash
        if current_question and current_options:
            self._save_question(current_question, current_options, correct_index)
        
        return self._validate_and_return()
    
    async def _parse_classic_format(self, paragraphs: list[str]) -> ParseResult:
        """
        Klassik format parse qilish:
        1. Savol matni?
        A) Variant
        *B) To'g'ri javob
        """
        current_question = None
        current_options = []
        correct_index = -1
        
        for para in paragraphs:
            # Savol ekanligini tekshirish
            question_match = self._match_question(para)
            if question_match:
                # Oldingi savolni saqlash
                if current_question and current_options:
                    self._save_question(current_question, current_options, correct_index)
                
                # Yangi savol
                current_question = question_match
                current_options = []
                correct_index = -1
                continue
            
            # Variant ekanligini tekshirish
            option_match = self._match_option(para)
            if option_match and current_question:
                option_text, is_correct = option_match
                current_options.append(option_text)
                if is_correct:
                    correct_index = len(current_options) - 1
        
        # Oxirgi savolni saqlash
        if current_question and current_options:
            self._save_question(current_question, current_options, correct_index)
        
        return self._validate_and_return()
    
    def _validate_and_return(self) -> ParseResult:
        """Natijani tekshirish va qaytarish"""
        # Natijani tekshirish
        if not self.questions:
            return ParseResult(
                success=False,
                questions=[],
                error_message="❌ Hech qanday savol topilmadi.\n\n"
                             "📝 Quyidagi formatlardan birini ishlating:\n\n"
                             "<b>Format 1:</b>\n"
                             "?Savol matni\n"
                             "+To'g'ri javob\n"
                             "=Noto'g'ri variant\n"
                             "=Noto'g'ri variant\n\n"
                             "<b>Format 2:</b>\n"
                             "1. Savol matni?\n"
                             "A) Birinchi variant\n"
                             "*B) To'g'ri javob\n"
                             "C) Uchinchi variant"
            )
        
        # Xatolarni tekshirish
        questions_without_correct = [
            i+1 for i, q in enumerate(self.questions) if q.correct_index == -1
        ]
        
        if questions_without_correct:
            return ParseResult(
                success=False,
                questions=[],
                error_message=f"❌ Quyidagi savollarda to'g'ri javob belgilanmagan:\n"
                             f"Savollar: {', '.join(map(str, questions_without_correct))}\n\n"
                             f"💡 To'g'ri javobni + belgisi bilan belgilang:\n"
                             f"+To'g'ri javob"
            )
        
        # Kam variantli savollar
        questions_few_options = [
            i+1 for i, q in enumerate(self.questions) if len(q.options) < 2
        ]
        
        if questions_few_options:
            return ParseResult(
                success=False,
                questions=[],
                error_message=f"❌ Quyidagi savollarda kamida 2 ta variant bo'lishi kerak:\n"
                             f"Savollar: {', '.join(map(str, questions_few_options))}"
            )
        
        return ParseResult(
            success=True,
            questions=self.questions,
            warnings=self.warnings
        )
    
    def _match_question(self, text: str) -> Optional[str]:
        """Savol ekanligini tekshirish (klassik format uchun)"""
        # ?bilan boshlanganlarni bu yerda skip qilamiz
        if text.startswith('?'):
            return None
            
        for pattern in self.QUESTION_PATTERNS[1:]:  # Birinchi pattern ? uchun
            match = re.match(pattern, text, re.IGNORECASE)
            if match:
                # Agar 2 ta guruh bo'lsa (raqam va matn)
                if len(match.groups()) >= 2:
                    return match.group(2).strip()
                else:
                    return match.group(1).strip()
        return None
    
    def _match_option(self, text: str) -> Optional[tuple[str, bool]]:
        """Variant ekanligini tekshirish. (matn, to'g'rimi) qaytaradi"""
        # + yoki = bilan boshlanganlarni skip qilamiz (yangi format)
        if text.startswith('+') or text.startswith('='):
            return None
        
        # To'g'ri javob belgisini tekshirish
        is_correct = False
        clean_text = text
        
        # Boshida * yoki + bor bo'lsa
        for marker in self.CORRECT_MARKERS:
            if text.startswith(marker):
                is_correct = True
                clean_text = text[1:].strip()
                break
        
        # Oxirida * yoki + bor bo'lsa
        if not is_correct:
            for marker in self.CORRECT_MARKERS:
                if text.endswith(marker):
                    is_correct = True
                    clean_text = text[:-1].strip()
                    break
        
        # Variant pattern'ini tekshirish
        for pattern in self.OPTION_PATTERNS:
            match = re.match(pattern, clean_text, re.IGNORECASE)
            if match:
                option_text = match.group(2).strip()
                # Oxiridagi * ni tozalash
                for marker in self.CORRECT_MARKERS:
                    if option_text.endswith(marker):
                        is_correct = True
                        option_text = option_text[:-1].strip()
                return (option_text, is_correct)
        
        return None
    
    def _save_question(self, question_text: str, options: list[str], correct_index: int):
        """Savolni saqlash"""
        question = Question(
            id=str(uuid.uuid4())[:8],
            text=question_text,
            options=options,
            correct_index=correct_index
        )
        self.questions.append(question)
